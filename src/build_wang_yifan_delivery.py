from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "outputs" / "wang_yifan_delivery"
OUT.mkdir(parents=True, exist_ok=True)
TARGET = OUT / "王伊梵_电力交易工作交付包_2026-08-06_v2_合同已签署.docx"

BLUE = "2E74B5"
DARK = "1F4D78"
LIGHT = "E8EEF5"
GRAY = "F2F4F7"
MUTED = RGBColor(90, 98, 108)


def set_font(run, size=11, bold=False, color=None):
    run.font.name = "Microsoft YaHei"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths_dxa[idx] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, text in enumerate(headers):
        cell = table.rows[0].cells[i]
        shade(cell, LIGHT)
        p = cell.paragraphs[0]
        set_font(p.add_run(text), 9.5, True, DARK)
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            set_font(p.add_run(str(text)), 9.2)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    set_font(p.add_run(text))


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    set_font(p.add_run(text))


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = section.bottom_margin = Inches(1)
section.left_margin = section.right_margin = Inches(1)
section.header_distance = section.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Microsoft YaHei"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25
for name, size, color, before, after in (
    ("Heading 1", 16, BLUE, 18, 10),
    ("Heading 2", 13, BLUE, 14, 7),
    ("Heading 3", 12, DARK, 10, 5),
):
    style = styles[name]
    style.font.name = "Microsoft YaHei"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)

header = section.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_font(header.add_run("电力交易算法工作包 | 内部讨论稿"), 9, False, "6B7280")
footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(footer.add_run("王伊梵 | 2026-08-06 | 输出仅供人工复核，不构成自动交易指令"), 8.5, False, "6B7280")

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(4)
set_font(p.add_run("工作交付包"), 24, True, DARK)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(14)
set_font(p.add_run("山东现货数据分析、价格预测、交易辅助算法与平台对接"), 14, False, BLUE)
for label, value in (("负责人", "王伊梵"), ("会议依据", "2026-07-22 电力交易讨论"), ("状态日期", "2026-08-06"), ("文档定位", "首次对齐会可直接使用的执行底稿")):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    set_font(p.add_run(f"{label}："), 10.5, True)
    set_font(p.add_run(value), 10.5)

doc.add_heading("1. 结论与任务状态", level=1)
p = doc.add_paragraph()
set_font(p.add_run("你的主责不是平台前端，而是算法侧闭环："), 11, True)
set_font(p.add_run("在统一数据口径上完成山东数据特征分析，开发价格/价差预测与交易辅助策略，固化为可回测、可陪跑、可嵌入平台的 SOP。"))

rows = [
    ("合同材料", "王伊梵", "已完成", "学校导师已完成签署；合同签署事项关闭"),
    ("Agent Box 学习", "王伊梵", "受阻", "缺少官网/账号/课程入口，无法替你登录验证"),
    ("山东数据基线分析", "王伊梵 + 陆璟行", "已形成基础", "已有 1-6 月价格、日用电量、分时用电量及分析脚本"),
    ("价格预测原型", "王伊梵", "已完成原型", "已有 6 月回测与 7 月预测；正式使用前需用最新数据滚动重训"),
    ("24 点 I/O 契约", "双方", "已起草", "随附 JSON 示例，字段和阈值须与陆璟行确认"),
    ("算法 SOP/复盘模板", "王伊梵", "已起草", "本工作包第 5、6 节"),
    ("亏损归因", "王伊梵", "数据受阻", "缺申报、成交、策略版本、人工调整、合同与结算数据"),
]
add_table(doc, ["工作项", "责任", "状态", "当前结论"], rows, [1800, 1350, 1150, 5060])

doc.add_heading("2. 已有数据与模型验收", level=1)
add_bullet(doc, "价格数据覆盖 2026-01-01 01:00 至 2026-07-01 00:00，共 4,344 个小时点，含日前与实时价格。")
add_bullet(doc, "用电数据包括日用电量和 24 点分时用电量；会议口径为 110 多家用户。公开输出必须仅使用稳定脱敏编号。")
add_bullet(doc, "现有预测原型使用岭回归（日历与滞后特征）、重复上周、小时-星期均值三个候选，并按回测择优或组合。")
add_bullet(doc, "6 月回测：日前组合模型 MAE 91.36 元/MWh、RMSE 119.78 元/MWh；实时岭回归 MAE 143.88 元/MWh、RMSE 184.16 元/MWh。")
add_bullet(doc, "风险识别仍偏弱：日前高价召回约 19.4%，实时高价召回为 0；该结果只能作为基线，不能据此自动下单。")
add_bullet(doc, "现有 7 月预测区间已经过期。下一次输出前必须追加最新真实价格、天气、负荷与市场状态，并按滚动窗口重新训练和回测。")

doc.add_heading("3. 山东数据分析提纲与特征清单", level=1)
add_table(doc, ["分析层", "必须输出", "建议特征"], [
    ("数据质量", "覆盖率、缺失、重复、异常、口径差异", "日期/时段完整性、单位、负值、极端值、跨表主键"),
    ("用户分群", "稳定型、波动型、峰谷显著型、异常型、重点运营型", "日均/峰值、变异系数、峰谷差、负荷率、工作日比、月趋势"),
    ("价格特征", "日前/实时分布、价差、负价、高价、异常时段", "小时、星期、月份、滞后 1/24/168h、滚动均值/波动"),
    ("负荷价格联动", "组合负荷与价格/价差的相关及分时差异", "组合负荷、预测偏差、峰谷位置、用户集中度"),
    ("模型准备", "训练/验证/回测切分及可用字段", "天气、新能源、机组、检修、阻塞、规则版本"),
], [1500, 3450, 4410])

doc.add_heading("4. 算法输入需求与平台契约", level=1)
p = doc.add_paragraph()
set_font(p.add_run("最小可运行输入："), 11, True)
set_font(p.add_run("市场日期、1-24 时段、日前/实时价格历史、组合与用户负荷历史、负荷预测、市场时区、模型版本。"))
p = doc.add_paragraph()
set_font(p.add_run("正式策略所需补充："), 11, True)
set_font(p.add_run("中长期持仓、零售义务、申报/成交、偏差考核、完整结算、天气、新能源预测、机组/检修/阻塞、人工调整原因及市场规则版本。"))
p = doc.add_paragraph()
set_font(p.add_run("接口草案文件："), 11, True)
set_font(p.add_run("outputs/wang_yifan_delivery/forecast_strategy_contract.example.json。双方需共同确认字段类型、空值策略、错误码、超时、版本兼容和审计 ID。"))

doc.add_heading("5. 预测与交易辅助 SOP（草案）", level=1)
for text in [
    "T-1 日收数并校验：检查 24 点完整性、单位、主键、异常值和数据更新时间；不合格则停止生成策略建议。",
    "锁定数据快照和版本：记录 request_id、数据区间、特征版本、模型版本、代码版本及市场规则版本。",
    "运行基线与候选模型：至少保留重复上周基线；滚动回测比较 MAE、RMSE、偏差、负价准确率与高价召回。",
    "生成分位数预测：逐时输出 P10/P50/P90、价差、负价/高价风险标志及原因码。",
    "施加策略约束：依据持仓、义务、最大电量、价格上限和最大日损失生成建议；关键输入为空时仅输出 HOLD。",
    "人工复核：交易员确认数据、异常时段、建议量价和风险提示；未经复核不得形成自动交易指令。",
    "盘后落库：保存预测、建议、人工调整、实际成交、实际价格、结算和偏差考核结果。",
    "每三天复盘：分解数据漂移、模型误差、策略偏差、执行偏差和收益影响；记录是否回滚或升级版本。",
]:
    add_number(doc, text)

doc.add_heading("6. 每三天复盘记录模板", level=1)
add_table(doc, ["字段", "填写内容"], [
    ("复盘周期 / 负责人", "____ 至 ____ / ____"),
    ("数据与模型版本", "数据快照 ____；特征 ____；模型 ____；规则 ____"),
    ("预测表现", "日前 MAE/RMSE/Bias ____；实时 MAE/RMSE/Bias ____；高价召回 ____"),
    ("策略与执行", "建议次数 ____；人工采纳/调整/拒绝 ____；主要原因 ____"),
    ("收益与风险", "收益 ____；偏差考核 ____；最大回撤/损失 ____；越界事件 ____"),
    ("问题归因", "数据 / 模型 / 策略 / 执行 / 市场变化：____"),
    ("后续动作", "动作 ____；责任人 ____；截止日期 ____；验收标准 ____"),
], [2300, 7060])

doc.add_heading("7. 不能在当前条件下代为完成的事项", level=1)
add_bullet(doc, "Agent Box 实操记录：需要官网/课程链接、可用账号或你允许使用的访问方式。")
add_bullet(doc, "与陆璟行的平台对齐：需要他的联系方式或由你转发本工作包和 JSON 草案。")
add_bullet(doc, "亏损归因和正式交易策略：必须补齐申报、成交、人工调整、合同、持仓、偏差考核和结算数据；缺失时不能给出可信的赚钱/亏损原因。")
add_bullet(doc, "当前日期后的预测：现有真实价格截至 2026-07-01，无法为 2026-08-06 之后生成可验证的现货预测。")

doc.add_heading("8. 你现在应立即推进的三件事", level=1)
for text in [
    "把 JSON 草案发给陆璟行，约 30 分钟对齐会，确认字段、接口、页面和人工复核状态机。",
    "索取 7 月至今价格/负荷及完整交易链路数据，并要求稳定脱敏 ID 与字段字典。",
    "提供 Agent Box 入口和账号；完成一个 24 点样例后记录输入、提示词、工具调用、输出和费用。",
]:
    add_number(doc, text)

doc.save(TARGET)
print(TARGET)
